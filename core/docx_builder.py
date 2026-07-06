# -*- coding: utf-8 -*-
"""
docx 조립 모듈. Call2~6 결과와 chart_generator가 만든 차트 이미지(BytesIO)를 받아
GUIDE_v2.1 섹션 6 스타일 규칙에 따라 최종 .docx를 BytesIO로 반환한다.
(Node.js 버전 build.js / 프로토타입 build_full.py의 python-docx 이식 검증을 그대로 재사용)
"""
import io
from docx import Document
from docx.shared import Twips, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PRIMARY = "1F4E79"
SECONDARY = "2E75B6"
LIGHT = "EBF3FA"
FONT_NAME = "Noto Sans CJK KR"  # 서브셋 폰트 내부 명칭과 일치해야 실제 렌더링에 반영됨

PAGE_W, PAGE_H = 11906, 16838
MARGIN = dict(top=1134, bottom=1134, left=1417, right=1134)
TABLE_WIDTH = PAGE_W - MARGIN["left"] - MARGIN["right"]


def _font(run, name=FONT_NAME, size=10.5, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)


def _shade(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def _width(cell, width_dxa):
    cell.width = Twips(width_dxa)
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(width_dxa)); tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


def _page_number_field(paragraph, field_type="PAGE"):
    run = paragraph.add_run()
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin')
    t = OxmlElement('w:instrText'); t.set(qn('xml:space'), 'preserve'); t.text = field_type
    e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end')
    run._r.append(b); run._r.append(t); run._r.append(e)
    _font(run, size=8, color="888888")


class DocxBuilder:
    def __init__(self, tech_name, purpose, scenario_label):
        self.doc = Document()
        self.tech_name = tech_name
        self.purpose = purpose
        self.scenario_label = scenario_label
        self._setup_page()
        self._setup_header_footer()

    def _setup_page(self):
        s = self.doc.sections[0]
        s.page_width = Twips(PAGE_W); s.page_height = Twips(PAGE_H)
        s.top_margin = Twips(MARGIN["top"]); s.bottom_margin = Twips(MARGIN["bottom"])
        s.left_margin = Twips(MARGIN["left"]); s.right_margin = Twips(MARGIN["right"])
        self.section = s

    def _setup_header_footer(self):
        hp = self.section.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _font(hp.add_run(f"{self.tech_name} | 기술동향 분석 보고서"), size=8, color="888888")

        fp = self.section.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _font(fp.add_run(f"{self.tech_name} | {self.purpose} | "), size=8, color="888888")
        _page_number_field(fp, "PAGE")
        _font(fp.add_run(" / "), size=8, color="888888")
        _page_number_field(fp, "NUMPAGES")

    def cover(self, date_str):
        for _ in range(6):
            self.doc.add_paragraph()
        p = self.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _font(p.add_run(f"{self.tech_name}\n기술동향 분석 보고서"), size=22, bold=True, color=PRIMARY)
        p2 = self.doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _font(p2.add_run("기술동향 분석 보고서"), size=14, color=SECONDARY)
        for _ in range(3):
            self.doc.add_paragraph()
        for label, val in [("작성 일자", date_str), ("분석 목적", self.purpose), ("적용 시나리오", self.scenario_label)]:
            pp = self.doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _font(pp.add_run(f"{label}: {val}"), size=11)
        disc = self.doc.add_paragraph(); disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        disc.paragraph_format.space_before = Pt(30)
        _font(disc.add_run("※ 본 보고서의 정량 수치는 별도 출처 표기가 없는 한 Claude 지식 기반 추정치(E)이며,\n"
                            "실제 의사결정 이전 원 데이터 재검증이 필요합니다."), size=9, italic=True, color="777777")

    def h1(self, text, page_break=True):
        if page_break:
            self.doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(12)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:space'), '6'); bottom.set(qn('w:color'), PRIMARY)
        pBdr.append(bottom); pPr.append(pBdr)
        _font(p.add_run(text), size=16, bold=True, color=PRIMARY)

    def h2(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(8)
        _font(p.add_run(text), size=13, bold=True, color=SECONDARY)

    def h3(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(6)
        _font(p.add_run(text), size=11, bold=True, color="333333")

    def body(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8); p.paragraph_format.line_spacing = 1.15
        _font(p.add_run(text), size=10.5)

    def bullet(self, text):
        p = self.doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(6)
        _font(p.add_run(text), size=10.5)

    def table(self, headers, rows, ratios):
        total = sum(ratios)
        widths = [round(r / total * TABLE_WIDTH) for r in ratios]
        widths[-1] += TABLE_WIDTH - sum(widths)
        t = self.doc.add_table(rows=1, cols=len(headers))
        t.autofit = False
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = t.rows[0].cells
        for i, h in enumerate(headers):
            _width(hdr[i], widths[i]); _shade(hdr[i], PRIMARY)
            pp = hdr[i].paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _font(pp.add_run(h), size=9.5, bold=True, color="FFFFFF")
        for ridx, row in enumerate(rows):
            cells = t.add_row().cells
            fill = LIGHT if ridx % 2 == 1 else None
            for i, val in enumerate(row):
                _width(cells[i], widths[i])
                if fill:
                    _shade(cells[i], fill)
                _font(cells[i].paragraphs[0].add_run(str(val)), size=9.5)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def image(self, buf: io.BytesIO, width_cm, caption):
        p = self.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        buf.seek(0)
        p.add_run().add_picture(buf, width=Cm(width_cm))
        cap = self.doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(10)
        _font(cap.add_run(caption), size=8.5, italic=True, color="555555")

    def note(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(10)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'FFF7E6')
        pPr.append(shd)
        pBdr = OxmlElement('w:pBdr')
        for side in ('top', 'bottom', 'left', 'right'):
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), '4'); el.set(qn('w:space'), '4'); el.set(qn('w:color'), 'EDA100')
            pBdr.append(el)
        pPr.append(pBdr)
        _font(p.add_run(f"※ {text}"), size=9.5, italic=True, color="6B4E00")

    def save_to_buffer(self) -> io.BytesIO:
        buf = io.BytesIO()
        self.doc.save(buf)
        buf.seek(0)
        return buf


def build_report_docx(tech_name, purpose, scenario_label, date_str, chapter_results, chart_images: dict) -> io.BytesIO:
    b = DocxBuilder(tech_name, purpose, scenario_label)
    b.cover(date_str)

    ch1 = chapter_results["call2"]["ch1"]; ch2 = chapter_results["call2"]["ch2"]
    ch3 = chapter_results["call3"]["ch3"]; ch4 = chapter_results["call3"]["ch4"]
    ch5 = chapter_results["call4"]["ch5"]
    ch6 = chapter_results["call5"]["ch6"]; ch7 = chapter_results["call5"]["ch7"]
    ch8 = chapter_results["call6"]["ch8"]; ch9 = chapter_results["call6"]["ch9"]

    # Ⅰ
    b.h1("Ⅰ. 기술 개요 및 배경", page_break=True)
    b.h2("1. 기술 정의 및 배경"); b.body(ch1["intro"]); b.body(ch1["background"])
    b.h2("2. 주요 성능 지표 (KPI)"); b.table(["지표", "설명"], ch1["kpi"], [1, 3])
    b.h2("3. 기술 분류 체계"); b.table(["분류", "범위"], ch1["classification_desc"], [1, 3])

    # Ⅱ
    b.h1("Ⅱ. 시장 환경 분석")
    b.h2("1. 시장 규모 및 성장 전망"); b.body(ch2["market_intro"])
    b.h2("2. 주요국 정책 동향"); b.table(["국가/지역", "정책 동향"], ch2["policy"], [1, 4])
    b.h2("3. 주요 기업 동향"); b.table(["기업명", "역할", "주요 동향"], ch2["players"], [1, 1, 3.2])
    b.h2("4. 시장 규모 전망 출처별 비교")
    b.image(chart_images["chart5_market"], 14.5, "[Chart 5] 출처별 시장 규모·CAGR 전망 비교 (E)")
    b.h2("5. 주요 기업 포지셔닝 맵")
    b.image(chart_images["chart6_positioning"], 13, "[Chart 6] 주요 기업 포지셔닝 맵 (E)")

    # Ⅲ
    b.h1("Ⅲ. 특허 정량 분석")
    b.h2("1. 분석 개요"); b.body(ch3["overview"]); b.table(["기술 분류", "추정 특허 건수"], ch3["counts"], [1, 2])
    b.h2("2. 전체 출원 동향")
    b.image(chart_images["chart4_country"], 15.5, "[Chart 4] 국가별 연도별 특허 출원 추이 및 누적 비중 (E)")
    b.h2("3. 주요 출원인 현황")
    b.table(["출원인", "건수 동향", "주요 국가", "핵심 기술 영역"], ch3["applicants"], [1, 1, 1, 2])
    b.h2("4. 주요 IPC 분류 동향"); b.table(["IPC 코드", "의미"], ch3["ipc"], [1, 3])

    # Ⅳ
    b.h1("Ⅳ. 핵심 기술 개요")
    b.table(["기술 분류", "핵심 원리", "기술 계보(진화 방향)", "응용/상용화 포인트"], ch4["summary"], [1, 2, 2.4, 1.8])

    # Ⅴ
    b.h1("Ⅴ. 주요 R&D 동향 분석")
    b.h2("1. R&D 분석 개요")
    b.image(chart_images["chart1_trend"], 15.5, "[Chart 1] 기술 영역별 연도별 논문 건수 추이 (E)")
    for t in ch5["trends"]:
        b.bullet(t)
    b.h2("2. 기술 영역별 R&D 동향")
    for label, key in [("A", "A"), ("B", "B"), ("C", "C"), ("D", "D")]:
        if key in ch5["by_area"]:
            b.h3(label); b.body(ch5["by_area"][key])
    b.h2("주요 연구 그룹 및 성과")
    b.table(["연구 그룹 유형(E)", "발표 매체", "핵심 성과 영역"], ch5["research_groups"], [1.6, 1.4, 2.4])
    b.h2("3. 핵심 키워드 트렌드")
    b.image(chart_images["chart2_keywords"], 15, "[Chart 2] 핵심 키워드 출현 빈도 비교 (E)")
    b.h2("4. R&D-특허 종합 교차 분석")
    b.image(chart_images["chart3_matrix"], 13, "[Chart 3] R&D-특허 포지셔닝 매트릭스 (E)")
    b.table(["기술 영역", "R&D 성장도", "특허 성숙도", "권장 전략"], ch5["cross_analysis"], [1.2, 1, 1, 2.6])

    # Ⅵ
    b.h1("Ⅵ. 기술 성장 단계 분석")
    b.table(["기술 분류", "성장 단계", "R&D 성장도", "시사점"], ch6["stages"], [1.2, 1.2, 1, 2.6])
    b.h2("종합 판단"); b.body(ch6["overall"])

    # Ⅶ
    b.h1("Ⅶ. 주요 출원인 IP 히스토리 분석")
    b.h2("1. 상위 출원인 IP 히스토리")
    b.table(["출원인", "시기별 핵심 출원 영역 변화(E)", "R&D-특허 전환 패턴"], ch7["history"], [1, 3.2, 2])
    b.h2("2. 의뢰 기관 보유 IP 현황 및 평가"); b.body(ch7["own_ip_note"])

    # Ⅷ
    b.h1("Ⅷ. 공백기술 도출 및 IP 포트폴리오 전략")
    b.h2("1. 공백기술 도출 (3P 분석: 특허·논문·시장 교차)")
    b.table(["No.", "공백기술명", "관련 선행특허", "신규 IP 창출안", "진입 가능성"], ch8["gaps"], [0.4, 1.8, 1.8, 2.4, 1.6])
    b.h2("2. IP 재구성 전략 (4관점)"); b.table(["관점", "전략 내용"], ch8["reorg_strategy"], [1, 3])

    # Ⅸ
    b.h1("Ⅸ. 결론 및 시사점")
    b.h2("1. 기술·R&D 동향 요약")
    for k in ch9["key_points"]:
        b.bullet(k)
    b.h2("2. 우선순위 행동 과제"); b.table(["시점", "과제명", "구체 행동"], ch9["tasks"], [1, 1.6, 3.4])
    b.h2("3. 한계 및 유의사항")
    for l in ch9["limitations"]:
        b.bullet(l)
    b.note("공백기술 및 IP 전략과 관련한 실제 출원 여부 판단은 반드시 전문 변리사의 선행기술조사(FTO) 및 정밀 검토를 거쳐야 한다.")

    return b.save_to_buffer()
